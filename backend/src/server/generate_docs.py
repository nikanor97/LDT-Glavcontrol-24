import json

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.server import properties


def set_cell_merge(cell, axis_type):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    if axis_type == 'start':
        vMerge = OxmlElement('w:vMerge')
        vMerge.set(qn('w:val'), 'restart')
        tcPr.append(vMerge)
    elif axis_type == 'continue':
        vMerge = OxmlElement('w:vMerge')
        vMerge.set(qn('w:val'), 'continue')
        tcPr.append(vMerge)

def generate_docx(items_list, item_properties, output_file):
    # Создаем новый документ
    doc = Document()

    # Первая страница с центровкой текста "Техническое задание"
    section = doc.sections[0]
    section.page_height = Pt(842)  # Высота страницы А4 в пунктах (высота)
    section.page_width = Pt(595.2)  # Ширина страницы А4 в пунктах (ширина)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Техническое задание')

    # Увеличение шрифта
    run.font.size = Pt(24)

    # Центрирование текста по ширине и высоте
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    section.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавляем пустые параграфы для центровки по высоте
    before_paragraph = doc.add_paragraph()
    before_paragraph.add_run('').add_break()
    before_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Центрация за счет разрывов параграфов
    for _ in range(15):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()  # Разрыв страницы

    # Вторую страницу создаем с таблицей
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'

    # Заполняем заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ п/п'
    hdr_cells[1].text = 'Наименование товара'
    hdr_cells[2].text = 'Наименование показателя'
    hdr_cells[3].text = 'Описание, значение'
    hdr_cells[4].text = 'Цена за единицу, руб.'
    hdr_cells[5].text = 'Количество'
    hdr_cells[6].text = 'Цена за все, руб.'

    # Заполняем таблицу данными из списка и словаря
    for i, item in enumerate(items_list, start=1):
        name = item.get("name", "")
        price = item.get("price", "")
        amount = item.get("amount", "")
        if name in item_properties:
            properties = item_properties[name].split(';')
            num_rows = len(properties)

            # Первая строка с объединением
            first_row_cells = table.add_row().cells
            first_row_cells[0].text = str(i)
            first_row_cells[1].text = name
            first_row_cells[2].text = properties[0]
            first_row_cells[3].text = ''
            first_row_cells[4].text = str(int(price))
            first_row_cells[5].text = str(int(amount))
            first_row_cells[6].text = str(int(price*amount))

            set_cell_merge(first_row_cells[0], 'start')
            set_cell_merge(first_row_cells[1], 'start')
            set_cell_merge(first_row_cells[4], 'start')
            set_cell_merge(first_row_cells[5], 'start')
            set_cell_merge(first_row_cells[6], 'start')
            # set_cell_merge(first_row_cells[3], 'start')

            for prop in properties[1:]:
                row_cells = table.add_row().cells
                row_cells[0].text = ''
                row_cells[1].text = ''
                row_cells[2].text = prop
                row_cells[3].text = ''
                row_cells[4].text = ''
                row_cells[5].text = ''
                row_cells[6].text = ''

                set_cell_merge(row_cells[0], 'continue')
                set_cell_merge(row_cells[1], 'continue')
                set_cell_merge(row_cells[4], 'continue')
                set_cell_merge(row_cells[5], 'continue')
                set_cell_merge(row_cells[6], 'continue')
                # set_cell_merge(row_cells[3], 'continue')

    # Сохраняем документ
    doc.save(output_file)

# Пример использования:
items_list = [
    {"name": "Краска штемпельная"}, 
    {"name": "Бумага для принтера"}
]

item_properties = {
    "Краска штемпельная": "Специальная штемпельная краска;Объем флакона;Страна происхождения;Артикул",
    "Бумага для принтера": "Формат;Плотность;Цвет;Производитель"
}


if __name__ == "__main__":

    with open("../outgoing/заявки/прогноз_год_заявка_1.json", 'r', encoding='utf-8') as file:
        order_1 = json.load(file)
    output_file = 'output.docx'
    generate_docx(order_1, properties.item_properties, output_file)
    print(f"Документ '{output_file}' был успешно создан.")



# <table>
#     <thead>
#         <tr>
#             <th>№ п/п</th>
#             <th>Наименование товара</th>
#             <th>Наименование показателя</th>
#             <th>Описание, значение</th>
#         </tr>
#     </thead>
#     <tbody>
#         <tr>
#             <td rowspan=2 align="center">1</td>
#             <td rowspan=2 align="center">Краска штемпельная</td>
#             <td align="center">Специальная штемпельная краска</td>
#             <td rowspan=1 align="center"></td>
#         </tr>
#         <tr>
#             <td align="center">Объем флакона</td>
#             <td rowspan=1 align="center"></td>
#         </tr>
#         <tr>
#            <td rowspan=1 align="center"></td>
#         </tr>
#     </tbody>
# </table>